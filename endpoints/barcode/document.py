from config import PDF_QR_SIZE, DOC_QR_SIZE
from app_logger import logger
from io import FileIO
import traceback

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import fitz


def add_doc_header_picture(doc_name: str, picture_name: str):
    try:
        document = Document(doc_name)
        sections = document.sections
        sections[0].different_first_page_header_footer = True
        header = sections[0].first_page_header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if len(header_paragraph.runs) > 0:
            header_run = header_paragraph.runs[-1].clear()
            header_run = header_paragraph.runs[0].clear()
        else:
            header_run = header_paragraph.add_run()
        header_run.add_picture(FileIO(picture_name, 'rb'), width=Cm(DOC_QR_SIZE))
        document.save(doc_name)
        return 0
    except Exception as err:
        lf = '\n'
        logger.error(f'{traceback.format_exc().replace(lf, "")}')
        return -1


def add_pdf_header_picture(doc_name: str, picture_name: str):
    try:
        doc = fitz.open(doc_name)
        img = open(picture_name, 'rb').read()
        img_xref = 0
        page = doc.load_page(0)
        rect = fitz.Rect(page.rect.width - PDF_QR_SIZE * 2, PDF_QR_SIZE, page.rect.width - PDF_QR_SIZE, PDF_QR_SIZE * 2)
        page.insert_image(rect, stream=img, xref=img_xref)
        doc.saveIncr()
        return 0
    except Exception as err:
        lf = '\n'
        logger.error(f'{traceback.format_exc().replace(lf, "")}')
        return -1


if __name__ == "__main__":
    pass
    #add_doc_header_picture('test2.docx', 'test.png')
    #add_pdf_header_picture('', '')
